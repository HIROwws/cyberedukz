from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build_docx():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Техническое задание\nCyberEdu KZ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Веб-приложение для обучения кибербезопасности, адаптированное под СНГ и Казахстан")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph("Версия: 1.0 | Формат проекта: MVP на 2 недели | Исполнитель: индивидуальная разработка")

    doc.add_heading("1. Назначение документа", level=1)
    doc.add_paragraph(
        "Документ описывает требования к MVP веб-приложения CyberEdu KZ: образовательной платформы "
        "по кибербезопасности с практическими заданиями, флагами, прогрессом, курсами и CTF-механиками. "
        "ТЗ подготовлено так, чтобы его можно было использовать как основу проекта в Jira."
    )

    doc.add_heading("2. Видение продукта", level=1)
    doc.add_paragraph(
        "CyberEdu KZ помогает студентам, начинающим специалистам и преподавателям изучать кибербезопасность "
        "на русском языке с локальным контекстом Казахстана и СНГ. По духу продукт похож на TryHackMe, "
        "Hack The Box Academy, PortSwigger Web Security Academy, picoCTF и OverTheWire, но MVP фокусируется "
        "на управляемом обучении, понятной навигации и локализованном контенте."
    )
    add_bullets(doc, [
        "Практико-ориентированное обучение: теория сразу связана с заданиями.",
        "Локализация: русский интерфейс, примеры для Казахстана, термины из местной образовательной среды.",
        "Доступность для новичков: маршруты обучения от основ Linux, сетей и веб-безопасности.",
        "Соревновательность: очки, флаги, рейтинг и прогресс по направлениям.",
    ])

    doc.add_heading("3. Цели MVP", level=1)
    add_bullets(doc, [
        "Показать рабочий прототип платформы с авторизацией, каталогом курсов, заданиями и проверкой флагов.",
        "Создать минимум 3 учебных направления: Linux/Networks, Web Security, Blue Team Basics.",
        "Реализовать личный кабинет с прогрессом, баллами и историей решенных заданий.",
        "Подготовить демонстрационный набор контента из 10-15 заданий без опасной эксплуатации реальных систем.",
        "Оформить проект в Jira: epics, stories, tasks, sprint board и критерии приемки.",
    ])

    doc.add_heading("4. Целевая аудитория", level=1)
    add_table(doc, ["Сегмент", "Потребность", "MVP-ценность"], [
        ["Студенты колледжей и вузов", "Понятный вход в кибербезопасность", "Маршруты, теория, простые практики"],
        ["Начинающие SOC-аналитики", "Базовые навыки анализа событий", "Blue Team задания, логи, инциденты"],
        ["Начинающие pentester'ы", "Безопасная практика веб-уязвимостей", "SQLi, XSS, auth flaws в песочнице"],
        ["Преподаватели", "Готовая среда для лабораторных", "Каталог заданий и прогресс студентов"],
        ["Казахстанская аудитория", "Локальный язык и контекст", "Русский UI, KZ-примеры, часовой пояс, валюта KZT в будущем"],
    ], widths=[1.5, 2.4, 2.4])

    doc.add_heading("5. Роли пользователей", level=1)
    add_table(doc, ["Роль", "Права в MVP"], [
        ["Гость", "Просмотр публичной страницы, регистрация, вход"],
        ["Студент", "Прохождение курсов, отправка флагов, просмотр прогресса и рейтинга"],
        ["Администратор/Автор", "Создание и редактирование курсов, модулей и заданий"],
    ], widths=[1.8, 4.7])

    doc.add_heading("6. Функциональные требования", level=1)
    add_table(doc, ["ID", "Требование", "Приоритет", "Критерий приемки"], [
        ["FR-1", "Регистрация, вход, выход и базовый профиль пользователя", "Must", "Пользователь может создать аккаунт, войти и выйти"],
        ["FR-2", "Каталог курсов с фильтрами по направлению и сложности", "Must", "Курсы отображаются карточками и открываются в детальную страницу"],
        ["FR-3", "Страница курса с модулями, уроками и заданиями", "Must", "Пользователь видит структуру курса и статус прохождения"],
        ["FR-4", "Практическое задание с описанием, подсказками и полем ввода флага", "Must", "Верный флаг засчитывается, неверный показывает ошибку"],
        ["FR-5", "Начисление очков и сохранение прогресса", "Must", "После решения задания меняются баллы и процент курса"],
        ["FR-6", "Личный кабинет", "Must", "Отображаются курсы, решенные задания, очки и уровень"],
        ["FR-7", "Рейтинг пользователей", "Should", "Список пользователей отсортирован по баллам"],
        ["FR-8", "Админ-панель для контента", "Should", "Администратор может создать курс, модуль и задание"],
        ["FR-9", "Локализация интерфейса на русский", "Must", "Основные экраны и ошибки отображаются на русском языке"],
        ["FR-10", "Безопасная модель флагов", "Must", "Флаги не хранятся в открытом виде в клиентском коде"],
    ], widths=[0.7, 2.5, 0.8, 2.5])

    doc.add_heading("7. Учебный контент MVP", level=1)
    add_table(doc, ["Направление", "Модули", "Примеры заданий"], [
        ["Linux & Networks", "Linux basics, TCP/IP, DNS, HTTP", "Найти файл, прочитать логи, определить порт, разобрать HTTP-запрос"],
        ["Web Security", "SQL Injection, XSS, Auth, Command Injection", "Обойти простой login, найти reflected XSS, понять insecure direct object reference"],
        ["Blue Team Basics", "Логи, SIEM-мышление, incident response", "Найти подозрительный IP, классифицировать событие, собрать таймлайн"],
    ], widths=[1.5, 2.1, 2.9])

    doc.add_heading("8. Нефункциональные требования", level=1)
    add_bullets(doc, [
        "Безопасность: пароли хранятся только в хэшированном виде; доступ к админке ограничен ролью.",
        "Производительность: основные страницы открываются менее чем за 2 секунды на локальном стенде.",
        "Адаптивность: интерфейс корректно работает на ноутбуке и мобильном экране.",
        "Поддерживаемость: код разделен на frontend, backend/API и слой данных.",
        "Контентная безопасность: задания не должны давать пользователю доступ к реальной инфраструктуре.",
        "Язык: интерфейс MVP на русском; английские термины допускаются как профессиональные названия уязвимостей.",
    ])

    doc.add_heading("9. Рекомендуемая архитектура MVP", level=1)
    doc.add_paragraph(
        "Для одиночной разработки за две недели рекомендуется стек, который быстро дает результат: "
        "Next.js или React + Node.js API, PostgreSQL/SQLite для данных, Prisma/ORM для моделей, "
        "JWT/session auth, Tailwind или готовая UI-библиотека для интерфейса. Если проект учебный и сроки жесткие, "
        "можно начать с SQLite, а затем перенести на PostgreSQL."
    )
    add_table(doc, ["Слой", "Компоненты"], [
        ["Frontend", "Главная, каталог курсов, страница курса, задание, профиль, рейтинг, админка"],
        ["Backend/API", "Auth, courses, modules, tasks, flag submission, progress, leaderboard"],
        ["Database", "Users, courses, modules, lessons, challenges, submissions, progress"],
        ["Security", "Password hashing, role checks, server-side flag validation, input validation"],
    ], widths=[1.4, 5.1])

    doc.add_heading("10. Модель данных", level=1)
    add_table(doc, ["Сущность", "Ключевые поля"], [
        ["User", "id, email, passwordHash, name, role, points, createdAt"],
        ["Course", "id, title, slug, description, track, difficulty, language, isPublished"],
        ["Module", "id, courseId, title, order"],
        ["Lesson", "id, moduleId, title, theoryMarkdown, order"],
        ["Challenge", "id, lessonId, title, description, flagHash, points, difficulty, hint"],
        ["Submission", "id, userId, challengeId, submittedValue, isCorrect, createdAt"],
        ["Progress", "id, userId, courseId, completedLessons, completedChallenges, percent"],
    ], widths=[1.5, 5.0])

    doc.add_heading("11. Основные пользовательские сценарии", level=1)
    add_numbers(doc, [
        "Пользователь регистрируется и попадает в каталог курсов.",
        "Пользователь выбирает курс Web Security Basics и открывает первый модуль.",
        "Пользователь читает краткую теорию и переходит к практическому заданию.",
        "Пользователь вводит флаг; сервер проверяет его и сохраняет результат.",
        "Пользователь видит обновленные очки, прогресс курса и место в рейтинге.",
        "Администратор создает новый курс, добавляет модуль и публикует задание.",
    ])

    doc.add_heading("12. Ограничения MVP", level=1)
    add_bullets(doc, [
        "MVP не включает полноценные виртуальные машины и изолированные лабораторные окружения.",
        "Практика строится на статических заданиях, мини-сценариях, безопасных HTML/API-демо и проверке флагов.",
        "Оплата, сертификаты, группы учебных заведений и интеграция с LMS выносятся за пределы MVP.",
        "Контент должен быть образовательным и не должен поощрять атаки на реальные системы.",
    ])

    doc.add_heading("13. План реализации на 14 дней", level=1)
    add_table(doc, ["Дни", "Фокус", "Результат"], [
        ["1", "Jira, дизайн MVP, репозиторий, структура данных", "Backlog, board, стартовый проект"],
        ["2-3", "Auth и базовый layout", "Регистрация, вход, навигация"],
        ["4-5", "Каталог курсов и страница курса", "Карточки, фильтры, модули"],
        ["6-7", "Задания и проверка флагов", "Challenge page, submissions, points"],
        ["8", "Профиль и прогресс", "Dashboard пользователя"],
        ["9", "Рейтинг", "Leaderboard"],
        ["10-11", "Админка контента", "CRUD курсов, модулей, заданий"],
        ["12", "Контент MVP", "10-15 заданий, 3 направления"],
        ["13", "Тестирование и исправления", "Проверенные сценарии, seed data"],
        ["14", "Демо, документация, финальная защита", "README, скриншоты, презентационный сценарий"],
    ], widths=[0.8, 2.6, 3.1])

    doc.add_heading("14. Definition of Done", level=1)
    add_bullets(doc, [
        "Код запущен локально без ручных правок конфигурации после README.",
        "Пользователь может пройти минимум один полный курс от входа до зачета флага.",
        "В базе есть демонстрационные пользователи, курсы и задания.",
        "Все критичные user stories в Jira закрыты и имеют критерии приемки.",
        "Подготовлены скриншоты или короткий демо-сценарий для защиты проекта.",
    ])

    doc.add_heading("15. Риски и меры", level=1)
    add_table(doc, ["Риск", "Вероятность", "Мера"], [
        ["Слишком большой объем для одного разработчика", "Высокая", "Жестко держать MVP, не делать VM/lab infra"],
        ["Нехватка времени на контент", "Средняя", "Сначала сделать 10 коротких заданий, затем расширять"],
        ["Проблемы с безопасностью флагов", "Средняя", "Проверять флаги только на сервере, хранить хэши"],
        ["Сложность админки", "Средняя", "Сделать простой CRUD без сложного редактора"],
    ], widths=[2.4, 1.1, 3.0])

    doc.add_paragraph("Приложение: Jira backlog см. в файле jira_backlog_cyberedukz.csv.")

    path = OUT / "CyberEdu_KZ_Technical_Spec.docx"
    doc.save(path)
    return path


def build_jira_csv():
    rows = [
        ["CyberEdu KZ MVP", "Epic", "", "Highest", "CyberEdu KZ MVP", "", "Главный epic проекта: образовательная платформа по кибербезопасности для СНГ/Казахстана."],
        ["Auth and User Profile", "Epic", "", "High", "Auth and User Profile", "CyberEdu KZ MVP", "Регистрация, вход, роли и профиль пользователя."],
        ["Course Learning Experience", "Epic", "", "High", "Course Learning Experience", "CyberEdu KZ MVP", "Каталог курсов, уроки, задания, прогресс."],
        ["Gamification and Progress", "Epic", "", "Medium", "Gamification and Progress", "CyberEdu KZ MVP", "Очки, рейтинг, статусы прохождения."],
        ["Admin Content Management", "Epic", "", "Medium", "Admin Content Management", "CyberEdu KZ MVP", "CRUD для курсов, модулей и заданий."],
        ["Localization and Demo Content", "Epic", "", "High", "Localization and Demo Content", "CyberEdu KZ MVP", "Русский интерфейс, KZ/CIS контекст, учебные задания."],
        ["Set up repository and project structure", "Task", "CyberEdu KZ MVP", "Highest", "", "CyberEdu KZ MVP", "Создать проект, настроить scripts, env example, README, базовую структуру frontend/backend."],
        ["Create database schema", "Task", "CyberEdu KZ MVP", "Highest", "", "CyberEdu KZ MVP", "Описать User, Course, Module, Lesson, Challenge, Submission, Progress."],
        ["As a visitor, I can register an account", "Story", "Auth and User Profile", "Highest", "", "Auth and User Profile", "Acceptance: email/password validation; password hash; duplicate email handled."],
        ["As a user, I can log in and log out", "Story", "Auth and User Profile", "Highest", "", "Auth and User Profile", "Acceptance: valid credentials create session; logout clears session; invalid credentials show Russian error."],
        ["As a user, I can view my profile", "Story", "Auth and User Profile", "High", "", "Auth and User Profile", "Acceptance: profile shows name, email, points, completed challenges, active courses."],
        ["As a learner, I can browse courses", "Story", "Course Learning Experience", "Highest", "", "Course Learning Experience", "Acceptance: course cards show title, track, difficulty, progress; filters work."],
        ["As a learner, I can open a course page", "Story", "Course Learning Experience", "Highest", "", "Course Learning Experience", "Acceptance: modules and lessons shown in order; locked/unlocked state is clear."],
        ["As a learner, I can read a lesson", "Story", "Course Learning Experience", "High", "", "Course Learning Experience", "Acceptance: theory markdown renders; next challenge link visible."],
        ["As a learner, I can submit a flag", "Story", "Course Learning Experience", "Highest", "", "Course Learning Experience", "Acceptance: correct flag awards points once; wrong flag does not award points."],
        ["Implement server-side flag validation", "Task", "Course Learning Experience", "Highest", "", "Course Learning Experience", "Compare submitted value with server-side hash; never expose valid flags in client bundle."],
        ["As a learner, I can see course progress", "Story", "Gamification and Progress", "High", "", "Gamification and Progress", "Acceptance: progress percent changes after completed tasks."],
        ["As a learner, I can see leaderboard", "Story", "Gamification and Progress", "Medium", "", "Gamification and Progress", "Acceptance: users sorted by points; current user visible."],
        ["As an admin, I can create and edit courses", "Story", "Admin Content Management", "Medium", "", "Admin Content Management", "Acceptance: admin-only form creates course with title, track, difficulty, description."],
        ["As an admin, I can create challenges", "Story", "Admin Content Management", "Medium", "", "Admin Content Management", "Acceptance: admin can add description, hint, points and flag; flag stored as hash."],
        ["Seed MVP demo content", "Task", "Localization and Demo Content", "Highest", "", "Localization and Demo Content", "Create 3 tracks and 10-15 starter tasks with safe educational content."],
        ["Localize core interface to Russian", "Task", "Localization and Demo Content", "High", "", "Localization and Demo Content", "Navigation, forms, errors and empty states are in Russian."],
        ["Add Kazakhstan/CIS context to content", "Task", "Localization and Demo Content", "Medium", "", "Localization and Demo Content", "Use local examples: universities, SOC roles, KZ timezone, regional terminology."],
        ["Write README and demo script", "Task", "CyberEdu KZ MVP", "High", "", "CyberEdu KZ MVP", "README explains setup; demo script describes 5-minute walkthrough."],
        ["Test critical user journey", "Task", "CyberEdu KZ MVP", "Highest", "", "CyberEdu KZ MVP", "Test registration -> course -> lesson -> flag -> progress -> leaderboard."],
    ]
    path = OUT / "jira_backlog_cyberedukz.csv"
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["Summary", "Issue Type", "Epic Link", "Priority", "Epic Name", "Component", "Description"])
        writer.writerows(rows)
    return path


def build_plan_md():
    text = """# CyberEdu KZ: 14-дневный план реализации

## Цель
За 14 дней собрать демонстрационный MVP веб-приложения для обучения кибербезопасности: регистрация, каталог курсов, уроки, задания с флагами, прогресс, рейтинг и простая админка контента.

## Рекомендуемая настройка Jira
- Project type: Scrum или Kanban. Для двух недель удобнее Scrum с одним Sprint `MVP Sprint 1`.
- Epics: `Auth and User Profile`, `Course Learning Experience`, `Gamification and Progress`, `Admin Content Management`, `Localization and Demo Content`.
- Issue types: Epic, Story, Task, Bug.
- Workflow: To Do -> In Progress -> Review/Test -> Done.
- Definition of Done: задача реализована, проверена локально, имеет понятный результат для демонстрации.

## Приоритеты
1. Сначала пользовательский путь: регистрация -> курс -> урок -> флаг -> прогресс.
2. Затем контент и локализация.
3. Потом админка, рейтинг и визуальная полировка.

## Ежедневный план
| День | Работа | Готовый результат |
|---|---|---|
| 1 | Jira, репозиторий, стек, схема БД | Backlog импортирован, проект запускается |
| 2 | Auth UI/API | Регистрация и вход |
| 3 | Профиль и роли | Студент/админ, защищенные страницы |
| 4 | Каталог курсов | Карточки, фильтры, seed courses |
| 5 | Страница курса | Модули, уроки, статусы |
| 6 | Страница задания | Описание, подсказка, ввод флага |
| 7 | Проверка флага | Server-side validation, submissions |
| 8 | Прогресс | Очки, проценты, completed state |
| 9 | Рейтинг | Leaderboard |
| 10 | Админка курсов | CRUD курса/модуля |
| 11 | Админка заданий | CRUD challenge, hash flag |
| 12 | Контент | 10-15 заданий, 3 направления |
| 13 | Тестирование | Исправления, пустые состояния, ошибки |
| 14 | Защита | README, скриншоты, demo script |

## MVP-контент
- Linux & Networks: 4-5 заданий.
- Web Security: 4-5 заданий.
- Blue Team Basics: 3-5 заданий.

## Что не делать в MVP
- Не поднимать реальные атакуемые виртуальные машины.
- Не делать платежи, сертификаты, группы университетов и сложную LMS.
- Не делать сложный редактор контента вместо простых форм.
"""
    path = OUT / "CyberEdu_KZ_14_day_plan.md"
    path.write_text(text, encoding="utf-8")
    return path


if __name__ == "__main__":
    for artifact in [build_docx(), build_jira_csv(), build_plan_md()]:
        print(artifact)
